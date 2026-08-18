#!/usr/bin/env python3
"""Safely add one dated record to existing employee-training DOCX tables.

The source directory is never modified.  The script copies its complete tree to a
new staging directory, updates only roster-matched DOCX files, and writes a report
outside the staged tree.  It intentionally stops for uncertain person, document,
table, header, date, or formatting matches.
"""

from __future__ import annotations

import argparse
import csv
from copy import deepcopy
from dataclasses import dataclass
from datetime import date
import json
from pathlib import Path
import shutil
import sys
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, Table, _Row
from lxml import etree


DEFAULT_TABLE_LABEL = "专项安全培训教育记录"
EXPECTED_HEADERS = ("年", "月", "日", "主要内容", "课时", "考核成绩")


class UpdateError(RuntimeError):
    """An input is not safe or specific enough to update."""


@dataclass(frozen=True)
class Record:
    year: int
    month: int
    day: int
    content: str
    hours: str

    @property
    def values(self) -> tuple[str, str, str, str, str]:
        return (str(self.year), str(self.month), str(self.day), self.content, self.hours)

    @property
    def date_value(self) -> date:
        return date(self.year, self.month, self.day)


@dataclass(frozen=True)
class PlannedUpdate:
    name: str
    result: str
    source_path: Path


def normalise(value: object) -> str:
    return "".join(str(value or "").replace("\u3000", " ").split())


def cell_text(cell: _Cell) -> str:
    return normalise(cell.text)


def row_values(row: _Row, count: int = 6) -> tuple[str, ...]:
    return tuple(cell_text(cell) for cell in row.cells[:count])


def row_is_blank(row: _Row) -> bool:
    return not any(cell_text(cell) for cell in row.cells)


def table_matches_label(table: Table, label: str) -> bool:
    wanted = normalise(label)
    return any(wanted in cell_text(cell) for row in table.rows for cell in row.cells)


def find_header_index(table: Table) -> int | None:
    expected = tuple(normalise(item) for item in EXPECTED_HEADERS)
    for index, row in enumerate(table.rows):
        if len(row.cells) >= len(expected) and row_values(row) == expected:
            return index
    return None


def matching_tables(document: Document, label: str) -> list[tuple[Table, int]]:
    matches = []
    for table in document.tables:
        if not table_matches_label(table, label):
            continue
        header_index = find_header_index(table)
        if header_index is not None:
            matches.append((table, header_index))
    return matches


def parse_row_date(row: _Row) -> date | None:
    values = row_values(row, 3)
    if not all(values):
        return None
    try:
        return date(int(values[0]), int(values[1]), int(values[2]))
    except ValueError:
        return None


def populated_data_rows(table: Table, header_index: int) -> list[tuple[int, _Row, date | None]]:
    return [
        (index, row, parse_row_date(row))
        for index, row in enumerate(table.rows[header_index + 1 :], start=header_index + 1)
        if not row_is_blank(row)
    ]


def find_existing_record(
    data_rows: Iterable[tuple[int, _Row, date | None]], record: Record
) -> tuple[int, _Row, date | None] | None:
    wanted = tuple(normalise(value) for value in record.values)
    for item in data_rows:
        row = item[1]
        if row_values(row, 5) == wanted:
            return item
    return None


def choose_insert_after(
    table: Table, header_index: int, record: Record
) -> tuple[_Row, _Row]:
    """Return (row after which to insert, row whose formatting is copied)."""
    data_rows = populated_data_rows(table, header_index)
    if not data_rows:
        raise UpdateError("matching table has no populated data row to copy formatting from")

    dated_rows = [item for item in data_rows if item[2] is not None]
    if not dated_rows:
        raise UpdateError("matching table has no dated data row; cannot determine chronological position")
    if len(dated_rows) != len(data_rows):
        raise UpdateError("matching table has a populated row without a valid date")
    dated_values = [item[2] for item in dated_rows]
    if dated_values != sorted(dated_values):
        raise UpdateError("existing dated records are not in chronological order")

    after = table.rows[header_index]
    for _, row, row_date in dated_rows:
        if row_date is not None and row_date <= record.date_value:
            after = row

    if after._tr is table.rows[header_index]._tr:
        format_source = dated_rows[0][1]
    else:
        format_source = after
    return after, format_source


def primary_format_r_pr(paragraph) -> object | None:
    """Return the rPr of the run that best represents the cell's character formatting.

    Prefer the first run that carries visible text, then the first run that has
    an rPr, then the first run of any kind. This keeps the copied formatting
    stable when a cell contains leading empty runs (common in WPS/Word files).
    """
    text_r_pr = None
    fallback_r_pr = None
    for run in paragraph.iter(qn("w:r")):
        r_pr = run.rPr
        if r_pr is None:
            continue
        if text_r_pr is None and any((node.text or "").strip() for node in run.findall(qn("w:t"))):
            text_r_pr = r_pr
        if fallback_r_pr is None:
            fallback_r_pr = r_pr
        if text_r_pr is not None:
            break
    if text_r_pr is not None:
        return text_r_pr
    if fallback_r_pr is not None:
        return fallback_r_pr
    first_run = next(iter(paragraph.iter(qn("w:r"))), None)
    return first_run.rPr if first_run is not None else None


def primary_style_fingerprint(cell: _Cell) -> tuple[bytes, bytes, bytes]:
    """Capture the primary cell, paragraph, and run formatting without its text."""
    tc_pr = cell._tc.tcPr
    paragraphs = [child for child in cell._tc if child.tag == qn("w:p")]
    paragraph = paragraphs[0] if paragraphs else None
    p_pr = paragraph.pPr if paragraph is not None else None
    r_pr = primary_format_r_pr(paragraph) if paragraph is not None else None
    return tuple(
        etree.tostring(item, method="c14n", exclusive=True, with_comments=False)
        if item is not None
        else b""
        for item in (tc_pr, p_pr, r_pr)
    )


def replace_cell_text_preserving_format(cell: _Cell, value: str) -> None:
    """Replace cell content while keeping the copied cell/paragraph/run properties."""
    paragraphs = [child for child in cell._tc if child.tag == qn("w:p")]
    if not paragraphs:
        paragraph = OxmlElement("w:p")
        cell._tc.append(paragraph)
    else:
        paragraph = paragraphs[0]
        for extra in paragraphs[1:]:
            cell._tc.remove(extra)

    chosen_r_pr = primary_format_r_pr(paragraph)
    source_r_pr = deepcopy(chosen_r_pr) if chosen_r_pr is not None else None
    for child in list(paragraph):
        if child.tag != qn("w:pPr"):
            paragraph.remove(child)

    run = OxmlElement("w:r")
    if source_r_pr is not None:
        run.append(source_r_pr)
    text = OxmlElement("w:t")
    if value[:1].isspace() or value[-1:].isspace():
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = value
    run.append(text)
    paragraph.append(run)


def insert_record(table: Table, header_index: int, record: Record, result: str) -> None:
    after, format_source = choose_insert_after(table, header_index, record)
    new_tr = deepcopy(format_source._tr)
    new_cells = [_Cell(tc, table) for tc in new_tr.tc_lst]
    if len(new_cells) < len(EXPECTED_HEADERS):
        raise UpdateError("format source has fewer than six cells")

    values = (*record.values, result)
    for cell, value in zip(new_cells[:6], values, strict=True):
        replace_cell_text_preserving_format(cell, value)

    table._tbl.insert(table._tbl.index(after._tr) + 1, new_tr)


def read_roster(path: Path) -> dict[str, str]:
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as stream:
            reader = csv.DictReader(stream)
            headers = reader.fieldnames or []
            if "姓名" not in headers or "考核成绩" not in headers:
                raise UpdateError("roster CSV must have exact headers: 姓名,考核成绩")
            roster: dict[str, str] = {}
            for row_number, row in enumerate(reader, start=2):
                name = normalise(row.get("姓名"))
                result = normalise(row.get("考核成绩"))
                if not name or not result:
                    raise UpdateError(f"roster row {row_number} has an empty 姓名 or 考核成绩")
                if name in roster:
                    raise UpdateError(f"duplicate person in roster: {name}")
                roster[name] = result
    except OSError as error:
        raise UpdateError(f"cannot read roster: {error}") from error
    if not roster:
        raise UpdateError("roster is empty")
    return roster


def ensure_safe_source_tree(source_dir: Path) -> None:
    if not source_dir.is_dir():
        raise UpdateError(f"source directory does not exist: {source_dir}")
    links = [path for path in source_dir.rglob("*") if path.is_symlink()]
    if links:
        raise UpdateError(f"source tree contains symbolic link: {links[0].relative_to(source_dir)}")


def document_candidates(source_dir: Path, name: str) -> list[Path]:
    candidates = []
    for path in source_dir.rglob("*.docx"):
        if path.name.startswith("~$"):
            continue
        relative_text = path.relative_to(source_dir).as_posix()
        if name in relative_text:
            candidates.append(path)
    return sorted(candidates)


def preflight_person(source_dir: Path, name: str, label: str) -> PlannedUpdate:
    candidates = document_candidates(source_dir, name)
    matching_documents = []
    for candidate in candidates:
        try:
            document = Document(candidate)
        except Exception as error:  # python-docx reports several package exceptions
            raise UpdateError(f"cannot open DOCX for {name}: {candidate.name}: {error}") from error
        matches = matching_tables(document, label)
        if len(matches) == 1:
            table, header_index = matches[0]
            if not populated_data_rows(table, header_index):
                raise UpdateError(f"{name}: matching table has no populated training record")
            matching_documents.append(candidate)
        elif len(matches) > 1:
            raise UpdateError(f"{name}: more than one matching table in {candidate.name}")
    if len(matching_documents) != 1:
        rendered = ", ".join(path.relative_to(source_dir).as_posix() for path in matching_documents) or "none"
        raise UpdateError(f"{name}: expected exactly one matching training-record DOCX, found {rendered}")
    return PlannedUpdate(name=name, result="", source_path=matching_documents[0])


def find_target_row(table: Table, header_index: int, record: Record) -> _Row:
    candidates = [
        row
        for _, row, _ in populated_data_rows(table, header_index)
        if row_values(row, 5) == tuple(normalise(value) for value in record.values)
    ]
    if len(candidates) != 1:
        raise UpdateError("saved DOCX does not contain exactly one requested record")
    return candidates[0]


def verify_saved_document(
    path: Path, label: str, record: Record, result: str
) -> None:
    document = Document(path)
    matches = matching_tables(document, label)
    if len(matches) != 1:
        raise UpdateError("saved DOCX no longer has exactly one matching training table")
    table, header_index = matches[0]
    row = find_target_row(table, header_index, record)
    actual = row_values(row, 6)
    expected = (*tuple(normalise(value) for value in record.values), normalise(result))
    if actual != expected:
        raise UpdateError(f"saved record values differ: {actual!r}")
    target_index = next(index for index, item in enumerate(table.rows) if item._tr is row._tr)
    if target_index == header_index + 1:
        if target_index + 1 >= len(table.rows):
            raise UpdateError("saved record has no adjacent populated formatting reference")
        reference = table.rows[target_index + 1]
    else:
        reference = table.rows[target_index - 1]
    for index, (cell, reference_cell) in enumerate(zip(row.cells[:6], reference.cells[:6], strict=True), start=1):
        if primary_style_fingerprint(cell) != primary_style_fingerprint(reference_cell):
            raise UpdateError(f"saved record column {index} no longer matches the copied formatting")


def process_person(
    source_dir: Path, output_dir: Path, plan: PlannedUpdate, label: str, record: Record
) -> dict[str, str]:
    relative = plan.source_path.relative_to(source_dir)
    target_path = output_dir / relative
    document = Document(target_path)
    matches = matching_tables(document, label)
    if len(matches) != 1:
        raise UpdateError(f"{plan.name}: matching table changed after preflight")
    table, header_index = matches[0]
    data_rows = populated_data_rows(table, header_index)
    existing = find_existing_record(data_rows, record)
    if existing is not None:
        actual_result = row_values(existing[1], 6)[5]
        if actual_result != normalise(plan.result):
            raise UpdateError(
                f"{plan.name}: same date/content/hours already exists with a different result: {actual_result}"
            )
        return {
            "人员": plan.name,
            "来源文件": relative.as_posix(),
            "状态": "already_present",
            "备注": "请求的日期、内容、课时和考核成绩已存在，未重复写入",
        }

    insert_record(table, header_index, record, plan.result)
    document.save(target_path)
    verify_saved_document(target_path, label, record, plan.result)
    return {
        "人员": plan.name,
        "来源文件": relative.as_posix(),
        "状态": "updated",
        "备注": "已按日期插入，并按相邻已填写记录复制格式",
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source_dir", type=Path, help="complete extracted employee-archive tree")
    parser.add_argument("roster_csv", type=Path, help="UTF-8 CSV with 姓名,考核成绩")
    parser.add_argument("output_dir", type=Path, help="new staging directory; must not already exist")
    parser.add_argument("--year", required=True, type=int)
    parser.add_argument("--month", required=True, type=int)
    parser.add_argument("--day", required=True, type=int)
    parser.add_argument("--content", required=True)
    parser.add_argument("--hours", required=True)
    parser.add_argument("--table-label", default=DEFAULT_TABLE_LABEL)
    parser.add_argument("--report", type=Path, required=True, help="JSON report outside output_dir")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source_dir = args.source_dir.resolve()
    output_dir = args.output_dir.resolve()
    report_path = args.report.resolve()
    record = Record(args.year, args.month, args.day, normalise(args.content), normalise(args.hours))
    label = normalise(args.table_label)
    created_output = False
    try:
        ensure_safe_source_tree(source_dir)
        if output_dir.exists():
            raise UpdateError(f"output directory already exists: {output_dir}")
        if (
            output_dir == source_dir
            or output_dir in source_dir.parents
            or source_dir in output_dir.parents
        ):
            raise UpdateError("output directory must be outside the source tree")
        if report_path == output_dir or output_dir in report_path.parents:
            raise UpdateError("report path must be outside the output tree")
        if report_path == source_dir or source_dir in report_path.parents:
            raise UpdateError("report path must be outside the source tree")
        if report_path.exists():
            raise UpdateError(f"report path already exists: {report_path}")

        roster = read_roster(args.roster_csv)
        plans = []
        for name, result in roster.items():
            plan = preflight_person(source_dir, name, label)
            plans.append(PlannedUpdate(name=plan.name, result=result, source_path=plan.source_path))
        source_paths = [plan.source_path for plan in plans]
        if len(set(source_paths)) != len(source_paths):
            raise UpdateError("more than one roster name resolves to the same training-record DOCX")

        shutil.copytree(source_dir, output_dir, copy_function=shutil.copy2)
        created_output = True
        report_rows = [process_person(source_dir, output_dir, plan, label, record) for plan in plans]
        report = {
            "table_label": label,
            "record": {
                "年": record.year,
                "月": record.month,
                "日": record.day,
                "主要内容": record.content,
                "课时": record.hours,
            },
            "results": report_rows,
        }
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return 0
    except Exception as error:
        if created_output and output_dir.exists():
            shutil.rmtree(output_dir)
        print(f"training-record update error: {error}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
